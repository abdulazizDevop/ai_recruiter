# import openai
# import json
# import os
# from typing import Dict, List
# import PyPDF2
# import docx
# import re

# class AIService:
#     def __init__(self, api_key: str):
#         """OpenAI API bilan ishlash uchun servis"""
#         openai.api_key = api_key
#         self.model = "gpt-4"  # yoki "gpt-4-turbo"
        
#     def extract_text_from_pdf(self, file_path: str) -> str:
#         """PDF dan matn chiqarish"""
#         try:
#             text = ""
#             with open(file_path, 'rb') as file:
#                 pdf_reader = PyPDF2.PdfReader(file)
#                 for page in pdf_reader.pages:
#                     text += page.extract_text() + "\n"
#             return text.strip()
#         except Exception as e:
#             print(f"PDF o'qishda xatolik: {e}")
#             return ""
    
#     def extract_text_from_docx(self, file_path: str) -> str:
#         """DOCX dan matn chiqarish"""
#         try:
#             doc = docx.Document(file_path)
#             text = ""
#             for paragraph in doc.paragraphs:
#                 text += paragraph.text + "\n"
#             return text.strip()
#         except Exception as e:
#             print(f"DOCX o'qishda xatolik: {e}")
#             return ""
    
#     def extract_text_from_file(self, file_path: str) -> str:
#         """Fayl formatiga qarab matn chiqarish"""
#         file_extension = os.path.splitext(file_path)[1].lower()
        
#         if file_extension == '.pdf':
#             return self.extract_text_from_pdf(file_path)
#         elif file_extension in ['.docx', '.doc']:
#             return self.extract_text_from_docx(file_path)
#         else:
#             return ""
    
#     def analyze_resume(self, resume_text: str, criteria: Dict, vacancy_info: Dict) -> Dict:
#         """Resume ni tahlil qilish"""
        
#         # AI uchun prompt tayyorlash
#         prompt = self._create_analysis_prompt(resume_text, criteria, vacancy_info)
        
#         try:
#             response = openai.chat.completions.create(
#                 model=self.model,
#                 messages=[
#                     {"role": "system", "content": "Siz professional HR mutaxassisigiz. Resume ni tahlil qilib, vakansiya talablariga mosligini baholaysiz."},
#                     {"role": "user", "content": prompt}
#                 ],
#                 temperature=0.1,
#                 max_tokens=1500
#             )
            
#             analysis_text = response.choices[0].message.content
            
#             # JSON formatga o'girish
#             analysis_data = self._parse_analysis_result(analysis_text)
            
#             return analysis_data
            
#         except Exception as e:
#             print(f"AI tahlilida xatolik: {e}")
#             return {
#                 "match_percentage": 0,
#                 "analysis": "Tahlil qilishda xatolik yuz berdi",
#                 "strengths": [],
#                 "weaknesses": [],
#                 "recommendations": []
#             }
    
#     def _create_analysis_prompt(self, resume_text: str, criteria: Dict, vacancy_info: Dict) -> str:
#         """AI uchun prompt yaratish"""
        
#         prompt = f"""
# RESUME TAHLILI UCHUN VAZIFA:

# VAKANSIYA MA'LUMOTLARI:
# - Nomi: {vacancy_info.get('title', 'N/A')}
# - Tavsif: {vacancy_info.get('description', 'N/A')}
# - Talablar: {vacancy_info.get('requirements', 'N/A')}

# AI TAHLIL KRITERIYLARI:
# - Ko'nikmalar: {criteria.get('skills', 'N/A')}
# - Tajriba: {criteria.get('experience', 'N/A')} yil
# - Ta'lim: {criteria.get('education', 'N/A')}
# - Tillar: {criteria.get('languages', 'N/A')}
# - Qo'shimcha: {criteria.get('additional', 'N/A')}

# RESUME MATNI:
# {resume_text}

# TAHLIL QILING VA QUYIDAGI JSON FORMATDA JAVOB BERING:
# {{
#     "match_percentage": 75,
#     "analysis": "Umumiy tahlil matni",
#     "strengths": ["Kuchli tomonlar ro'yxati"],
#     "weaknesses": ["Zaif tomonlar ro'yxati"],
#     "recommendations": ["Tavsiyalar ro'yxati"],
#     "skills_match": {{"python": true, "javascript": false}},
#     "experience_match": true,
#     "education_match": true,
#     "languages_match": true
# }}

# ESLATMA: 
# - match_percentage 0 dan 100 gacha bo'lishi kerak
# - Har bir kriteriyani alohida baholang
# - 60% dan past bo'lsa, rad etish sabablari aniq ko'rsating
# """
        
#         return prompt
    
#     def _parse_analysis_result(self, analysis_text: str) -> Dict:
#         """AI javobini parse qilish"""
#         try:
#             # JSON qismini topish
#             json_start = analysis_text.find('{')
#             json_end = analysis_text.rfind('}') + 1
            
#             if json_start != -1 and json_end != -1:
#                 json_str = analysis_text[json_start:json_end]
#                 return json.loads(json_str)
#             else:
#                 # Agar JSON topilmasa, oddiy tahlil
#                 return {
#                     "match_percentage": 50,
#                     "analysis": analysis_text,
#                     "strengths": [],
#                     "weaknesses": [],
#                     "recommendations": []
#                 }
                
#         except json.JSONDecodeError:
#             return {
#                 "match_percentage": 50,
#                 "analysis": analysis_text,
#                 "strengths": [],
#                 "weaknesses": [],
#                 "recommendations": []
#             }
    
#     def evaluate_interview_answers(self, questions: List[str], answers: List[str], 
#                                  vacancy_info: Dict, criteria: Dict) -> Dict:
#         """Suhbat javoblarini baholash"""
        
#         # Savol-javoblarni birlashtirish
#         qa_pairs = []
#         for i, (q, a) in enumerate(zip(questions, answers), 1):
#             qa_pairs.append(f"Savol {i}: {q}\nJavob {i}: {a}")
        
#         qa_text = "\n\n".join(qa_pairs)
        
#         prompt = f"""
# SUHBAT JAVOBLARINI BAHOLASH:

# VAKANSIYA MA'LUMOTLARI:
# - Nomi: {vacancy_info.get('title', 'N/A')}
# - Talablar: {vacancy_info.get('requirements', 'N/A')}

# KRITERILAR:
# - Ko'nikmalar: {criteria.get('skills', 'N/A')}
# - Tajriba: {criteria.get('experience', 'N/A')} yil

# SAVOL-JAVOBLAR:
# {qa_text}

# JAVOBLARNI BAHOLANG VA QUYIDAGI JSON FORMATDA NATIJA BERING:
# {{
#     "total_score": 85,
#     "individual_scores": [80, 90, 85],
#     "evaluation": "Umumiy baholash matni",
#     "strengths": ["Kuchli javoblar"],
#     "improvements": ["Yaxshilanishi kerak bo'lgan tomonlar"],
#     "recommendation": "Qabul qilish yoki rad etish tavsiyasi",
#     "passed": true
# }}

# ESLATMA:
# - total_score 0 dan 100 gacha
# - individual_scores har bir savol uchun ball
# - passed: 70 ball va undan yuqori bo'lsa true
# - Javoblarning sifati, aniqlik va professionalligini baholang
# """
        
#         try:
#             response = openai.chat.completions.create(
#                 model=self.model,
#                 messages=[
#                     {"role": "system", "content": "Siz professional HR suhbat o'tkazuvchisigiz. Nomzod javoblarini baholaysiz."},
#                     {"role": "user", "content": prompt}
#                 ],
#                 temperature=0.1,
#                 max_tokens=1000
#             )
            
#             result_text = response.choices[0].message.content
            
#             # JSON parse qilish
#             try:
#                 json_start = result_text.find('{')
#                 json_end = result_text.rfind('}') + 1
#                 json_str = result_text[json_start:json_end]
#                 result = json.loads(json_str)
                
#                 # Natijani tekshirish
#                 if result.get('total_score', 0) >= 70:
#                     result['passed'] = True
#                 else:
#                     result['passed'] = False
                    
#                 return result
                
#             except json.JSONDecodeError:
#                 return {
#                     "total_score": 60,
#                     "evaluation": result_text,
#                     "passed": False,
#                     "individual_scores": [60] * len(questions),
#                     "strengths": [],
#                     "improvements": [],
#                     "recommendation": "Baholashda xatolik"
#                 }
                
#         except Exception as e:
#             print(f"Suhbat baholashda xatolik: {e}")
#             return {
#                 "total_score": 0,
#                 "evaluation": "Baholashda xatolik yuz berdi",
#                 "passed": False,
#                 "individual_scores": [0] * len(questions),
#                 "strengths": [],
#                 "improvements": [],
#                 "recommendation": "Qayta urinib ko'ring"
#             }
    
#     def generate_interview_questions(self, criteria: Dict, vacancy_info: Dict, count: int = 3) -> List[str]:
#         """Suhbat savollarini avtomatik yaratish (ixtiyoriy)"""
        
#         prompt = f"""
# SUHBAT SAVOLLARI YARATISH:

# VAKANSIYA: {vacancy_info.get('title', 'N/A')}
# TALABLAR: {vacancy_info.get('requirements', 'N/A')}
# KRITERILAR: Ko'nikmalar: {criteria.get('skills', 'N/A')}, 
# Tajriba: {criteria.get('experience', 'N/A')} yil

# {count} ta professional suhbat savolini yarating.
# Savollar vakansiya talablariga mos va nomzod qobiliyatini baholashga yo'naltirilgan bo'lsin.

# Javobni quyidagi formatda bering:
# 1. Savol matni
# 2. Savol matni
# 3. Savol matni
# """
        
#         try:
#             response = openai.chat.completions.create(
#                 model=self.model,
#                 messages=[
#                     {"role": "system", "content": "Siz professional HR mutaxassisigiz. Suhbat savollarini yaratasiz."},
#                     {"role": "user", "content": prompt}
#                 ],
#                 temperature=0.3,
#                 max_tokens=800
#             )
            
#             questions_text = response.choices[0].message.content
            
#             # Savollarni ajratish
#             questions = []
#             lines = questions_text.split('\n')
            
#             for line in lines:
#                 line = line.strip()
#                 # Raqam bilan boshlanuvchi satrlarni topish
#                 if re.match(r'^\d+\.', line):
#                     question = re.sub(r'^\d+\.\s*', '', line)
#                     if question:
#                         questions.append(question)
            
#             # Agar kerakli miqdorda savol topilamsa
#             if len(questions) >= count:
#                 return questions[:count]
#             else:
#                 # Fallback: standart savollar
#                 return [
#                     "O'zingiz haqingizda qisqacha gapiring.",
#                     "Nega aynan bu vakansiyani tanladingiz?",
#                     "Kelajakdagi rejalaringiz qanday?"
#                 ][:count]
                
#         except Exception as e:
#             print(f"Savollar yaratishda xatolik: {e}")
#             return [
#                 "O'zingiz haqingizda qisqacha gapiring.",
#                 "Nega aynan bu vakansiyani tanladingiz?",
#                 "Kelajakdagi rejalaringiz qanday?"
#             ][:count]

# # Yordamchi funksiyalar
# def validate_file_format(filename: str) -> bool:
#     """Fayl formatini tekshirish"""
#     allowed_extensions = ['.pdf', '.doc', '.docx']
#     file_extension = os.path.splitext(filename)[1].lower()
#     return file_extension in allowed_extensions

# def get_file_size_mb(file_path: str) -> float:
#     """Fayl hajmini MB da qaytarish"""
#     try:
#         size_bytes = os.path.getsize(file_path)
#         size_mb = size_bytes / (1024 * 1024)
#         return round(size_mb, 2)
#     except:
#         return 0

# def create_safe_filename(original_filename: str, user_id: int) -> str:
#     """Xavfsiz fayl nomi yaratish"""
#     import time
#     import re
    
#     # Faqat harflar, raqamlar va nuqtalarni qoldirish
#     safe_name = re.sub(r'[^a-zA-Z0-9.]', '_', original_filename)
    
#     # Vaqt va user_id qo'shish
#     timestamp = int(time.time())
#     name, extension = os.path.splitext(safe_name)
    
#     return f"{user_id}_{timestamp}_{name}{extension}"


# -*- coding: utf-8 -*-

import asyncio
import json
import logging
import re
import time
import os
from datetime import datetime
from typing import Dict, List, Optional, Any, Union, Tuple
from pathlib import Path

# Third-party imports
import openai
from openai import AsyncOpenAI
import PyPDF2
import docx
from docx import Document
import pymupdf  # PyMuPDF - PDF parsing uchun
import aiofiles
import aiohttp
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# Local imports
from config import Config

# Logger setup
logger = logging.getLogger(__name__)

# ===========================================
# AI SERVICE CLASS - Professional Version
# ===========================================

class AIService:
    """
    OpenAI API bilan professional integratsiya
    Resume tahlil qilish, suhbat baholash, AI assistant funksiyalari
    """
    
    def __init__(self, api_key: str, model: str = None, max_tokens: int = None, temperature: float = None):
        """AIService ni inicializatsiya qilish"""
        self.api_key = api_key
        self.model = model or Config.OPENAI_MODEL
        self.max_tokens = max_tokens or Config.OPENAI_MAX_TOKENS
        self.temperature = temperature or Config.OPENAI_TEMPERATURE
        
        # Async OpenAI client
        self.client = AsyncOpenAI(api_key=api_key)
        
        # Rate limiting
        self.last_request_time = 0
        self.request_count = 0
        self.request_window_start = time.time()
        
        logger.info(f"AI Service initialized with model: {self.model}")
    
    async def test_connection(self) -> bool:
        """OpenAI API connection ni test qilish"""
        try:
            response = await self.client.models.list()
            logger.info("OpenAI API connection successful")
            return True
        except Exception as e:
            logger.error(f"OpenAI API connection failed: {e}")
            return False
    
    # ===========================================
    # FILE PROCESSING METHODS
    # ===========================================
    
    async def extract_text_from_file(self, file_path: str) -> str:
        """Fayl formatiga qarab matn chiqarish"""
        try:
            file_extension = Path(file_path).suffix.lower()
            
            if file_extension == '.pdf':
                return await self._extract_text_from_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return await self._extract_text_from_docx(file_path)
            else:
                logger.warning(f"Unsupported file format: {file_extension}")
                return ""
                
        except Exception as e:
            logger.error(f"File text extraction error: {e}")
            return ""
    
    async def _extract_text_from_pdf(self, file_path: str) -> str:
        """PDF dan matn chiqarish - Professional version"""
        text = ""
        
        try:
            # PyMuPDF (fitz) ishlatish - yaxshiroq natija beradi
            if 'fitz' in globals():
                import fitz
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            else:
                # Fallback - PyPDF2
                async with aiofiles.open(file_path, 'rb') as file:
                    content = await file.read()
                    import io
                    pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
            
            return self._clean_extracted_text(text)
            
        except Exception as e:
            logger.error(f"PDF text extraction error: {e}")
            return ""
    
    async def _extract_text_from_docx(self, file_path: str) -> str:
        """DOCX dan matn chiqarish - Async version"""
        try:
            # DOCX faylni async tarzda o'qish
            doc = Document(file_path)
            text = ""
            
            # Paragraflar
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Jadvallar
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                text += "\n"
            
            return self._clean_extracted_text(text)
            
        except Exception as e:
            logger.error(f"DOCX text extraction error: {e}")
            return ""
    
    def _clean_extracted_text(self, text: str) -> str:
        """Chiqarilgan matnni tozalash"""
        if not text:
            return ""
        
        # Ko'p bo'shliqlarni bitta bo'shliqqa aylantirish
        text = re.sub(r'\s+', ' ', text)
        
        # Maxsus belgilarni tozalash
        text = re.sub(r'[^\w\s\-.,@():]', ' ', text)
        
        # Ko'p nuqtalarni tozalash
        text = re.sub(r'\.{3,}', '', text)
        
        return text.strip()
    
    # ===========================================
    # RESUME PARSING METHODS
    # ===========================================
    
    async def parse_resume(self, file_path: str) -> Optional[Dict[str, Any]]:
        """Resume ni to'liq parse qilish - AI + rule-based approach"""
        try:
            # Matnni chiqarish
            resume_text = await self.extract_text_from_file(file_path)
            if not resume_text:
                logger.warning("No text extracted from resume")
                return None
            
            # AI orqali strukturlangan data olish
            parsed_data = await self._ai_parse_resume(resume_text)
            
            # Rule-based parsing bilan qo'shimcha ma'lumot
            rule_based_data = self._rule_based_parsing(resume_text)
            
            # Ikki usulning natijalarini birlashtirish
            combined_data = self._combine_parsing_results(parsed_data, rule_based_data)
            combined_data['raw_text'] = resume_text
            combined_data['file_path'] = file_path
            combined_data['parsed_at'] = datetime.now().isoformat()
            
            logger.info(f"Resume parsed successfully: {len(resume_text)} characters")
            return combined_data
            
        except Exception as e:
            logger.error(f"Resume parsing error: {e}")
            return None
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=4, max=10),
        retry=retry_if_exception_type(openai.RateLimitError)
    )
    async def _ai_parse_resume(self, resume_text: str) -> Dict[str, Any]:
        """AI orqali resume ni strukturlangan formatga o'tkazish"""
        prompt = f"""
Quyidagi resume matnini tahlil qilib, JSON formatda strukturlangan ma'lumot chiqaring:

RESUME MATNI:
{resume_text[:4000]}  # Token limitini hisobga olish

Quyidagi JSON formatda javob bering:
{{
    "personal_info": {{
        "name": "To'liq ism",
        "email": "email@example.com", 
        "phone": "+998901234567",
        "address": "Manzil",
        "date_of_birth": "1990-01-01"
    }},
    "education": [
        {{
            "degree": "Bakalavr",
            "field": "Kompyuter fanlari", 
            "institution": "Universitet nomi",
            "graduation_year": 2020,
            "gpa": "3.8/4.0"
        }}
    ],
    "work_experience": [
        {{
            "position": "Lavozim nomi",
            "company": "Kompaniya nomi",
            "duration": "2020-2023",
            "responsibilities": ["Mas'uliyat 1", "Mas'uliyat 2"],
            "achievements": ["Yutuq 1", "Yutuq 2"]
        }}
    ],
    "skills": {{
        "technical_skills": ["Python", "JavaScript", "SQL"],
        "soft_skills": ["Liderlik", "Muloqot", "Muammolarni hal qilish"],
        "languages": ["O'zbek (Ona tili)", "Ingliz (B2)", "Rus (C1)"]
    }},
    "projects": [
        {{
            "name": "Loyiha nomi",
            "description": "Qisqacha tavsif",
            "technologies": ["React", "Node.js"],
            "link": "https://github.com/..."
        }}
    ],
    "certifications": [
        {{
            "name": "Sertifikat nomi",
            "organization": "Beruvchi tashkilot",
            "date": "2023-01-01",
            "credential_id": "12345"
        }}
    ],
    "summary": "Nomzod haqida umumiy xulosalar"
}}

ESLATMA: Agar biror ma'lumot topilmasa, null yoki bo'sh massiv qaytaring.
"""
        
        try:
            response = await self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "system", 
                        "content": "Siz professional HR mutaxassisigiz. Resume larni to'g'ri tahlil qilib, strukturlangan ma'lumot chiqarasiz."
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=2500
            )
            
            result_text = response.choices[0].message.content
            
            # JSON ni parse qilish
            json_match = re.search(r'\{.*\}', result_text, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                logger.warning("AI response contains no valid JSON")
                return self._get_empty_resume_structure()
                
        except json.JSONDecodeError as e:
            logger.error(f"AI resume parsing JSON error: {e}")
            return self._get_empty_resume_structure()
        except Exception as e:
            logger.error(f"AI resume parsing error: {e}")
            return self._get_empty_resume_structure()
    
    def _rule_based_parsing(self, text: str) -> Dict[str, Any]:
        """Rule-based parsing - Universal approach"""
        result = {}
        
        try:
            # Email topish
            email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
            emails = re.findall(email_pattern, text)
            result['emails'] = emails[:3]
            
            # Telefon raqam topish
            phone_patterns = [
                r'\+998\d{9}',
                r'\+\d{1,3}\s?\d{3,4}\s?\d{3,4}\s?\d{3,4}',
                r'\d{2,3}-\d{3,4}-\d{3,4}',
            ]
            phones = []
            for pattern in phone_patterns:
                phones.extend(re.findall(pattern, text))
            result['phones'] = phones[:3]
            
            # Yillar topish
            year_pattern = r'\b(19|20)\d{2}\b'
            years = sorted(set(re.findall(year_pattern, text)))
            result['years'] = years
            
            # ========================================
            # UNIVERSAL SKILLS DETECTION - AI ga topshirish
            # ========================================
            
            # Rule-based da faqat umumiy pattern topish
            # Skilllar AI orqali domain-specific holda topiladi
            
            # Sertifikat pattern
            cert_patterns = [
                r'certified?\s+in\s+([A-Za-z\s,]+)',
                r'certification:\s*([A-Za-z\s,]+)',
                r'certificate\s+of\s+([A-Za-z\s,]+)'
            ]
            
            certifications = []
            for pattern in cert_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                certifications.extend(matches)
            
            result['detected_certifications'] = certifications[:5]
            
            # Software/Tools mention (universal approach)
            # Faqat umumiy pattern - specific skills AI ga
            software_patterns = [
                r'proficient\s+in\s+([A-Za-z0-9\s,.-]+)',
                r'experienced?\s+with\s+([A-Za-z0-9\s,.-]+)',
                r'skilled?\s+in\s+([A-Za-z0-9\s,.-]+)'
            ]
            
            software_mentions = []
            for pattern in software_patterns:
                matches = re.findall(pattern, text, re.IGNORECASE)
                software_mentions.extend(matches)
            
            result['software_mentions'] = software_mentions[:10]
            
            return result
            
        except Exception as e:
            logger.error(f"Rule-based parsing error: {e}")
            return {}

    # AI ga to'liq skill detection topshirish
    async def _ai_parse_resume(self, resume_text: str) -> Dict[str, Any]:
        """AI orqali resume ni strukturlangan formatga o'tkazish - Universal"""
        prompt = f"""
    Quyidagi resume matnini tahlil qilib, JSON formatda strukturlangan ma'lumot chiqaring.

    MUHIM: Resume qaysi sohaga tegishli bo'lsa, o'sha sohaga xos skilllar va terminlarni toping.
    - Agar IT resume bo'lsa: programming languages, frameworks, databases
    - Agar tibbiyot bo'lsa: medical procedures, certifications, specializations  
    - Agar moliya bo'lsa: financial analysis, accounting software, regulations
    - Agar ta'lim bo'lsa: teaching methods, curriculum, educational technologies
    - Va boshqa har qanday soha uchun ham

    RESUME MATNI:
    {resume_text[:5000]}

    JSON formatda javob bering:
    {{
        "personal_info": {{
            "name": "To'liq ism",
            "email": "email@example.com", 
            "phone": "+998901234567",
            "address": "Manzil"
        }},
        "education": [...],
        "work_experience": [...],
        "skills": {{
            "domain_specific_skills": ["Sohaga xos ko'nikmalar - avtomatik aniqlang"],
            "technical_skills": ["Texnik ko'nikmalar - universal"],
            "soft_skills": ["Yumshoq ko'nikmalar"],
            "languages": ["Tillar"],
            "certifications": ["Sertifikatlar - sohaga bog'liq"],
            "tools_software": ["Dasturlar va vositalar"]
        }},
        "industry_domain": "Aniqlangan soha nomi (IT, Healthcare, Finance, Education, etc)",
        "projects": [...],
        "summary": "Resume xulosasi"
    }}

    ESLATMA: 
    1. Resume matnidan sohani avtomatik aniqlang
    2. O'sha sohaga mos skilllar va terminlarni toping  
    3. Universal yondashuv - har qanday sohani qo'llab-quvvatlang
    4. Agar aniq soha aniqlanmasa, "General/Mixed" deb belgilang
    """
    
    def _get_empty_resume_structure(self) -> Dict[str, Any]:
        """Bo'sh resume strukturasini qaytarish"""
        return {
            "personal_info": {
                "name": None,
                "email": None,
                "phone": None,
                "address": None,
                "date_of_birth": None
            },
            "education": [],
            "work_experience": [],
            "skills": {
                "technical_skills": [],
                "soft_skills": [],
                "languages": []
            },
            "projects": [],
            "certifications": [],
            "summary": None
        }
    
    # ===========================================
    # COMPATIBILITY ANALYSIS
    # ===========================================
    
    async def analyze_resume_compatibility(self, parsed_resume: Dict[str, Any], 
                                         ai_criteria: Dict[str, Any], 
                                         ai_prompt: str) -> Dict[str, Any]:
        """Resume va vakansiya kriteriylari orasidagi moslikni tahlil qilish"""
        try:
            # Base prompt yaratish
            analysis_prompt = self._create_compatibility_prompt(
                parsed_resume, ai_criteria, ai_prompt
            )
            
            response = await self._make_ai_request(
                system_message="Siz professional HR mutaxassisigiz. Resume va vakansiya kriteriylari orasidagi moslikni aniq baholaysiz.",
                user_message=analysis_prompt,
                max_tokens=1500
            )
            
            # Natijani parse qilish
            analysis_result = self._parse_compatibility_result(response)
            
            # Qo'shimcha validatsiya
            analysis_result = self._validate_compatibility_result(analysis_result)
            
            logger.info(f"Compatibility analysis completed: {analysis_result.get('compatibility_score', 0)}%")
            return analysis_result
            
        except Exception as e:
            logger.error(f"Compatibility analysis error: {e}")
            return self._get_default_compatibility_result()
    
    def _create_compatibility_prompt(self, resume: Dict, criteria: Dict, custom_prompt: str) -> str:
        """Moslik tahlili uchun prompt yaratish"""
        
        # Resume ma'lumotlarini formatlash
        resume_summary = self._format_resume_for_analysis(resume)
        
        prompt = f"""
RESUME VA VAKANSIYA MOSLIGINI TAHLIL QILING:

RESUME MA'LUMOTLARI:
{resume_summary}

VAKANSIYA KRITERIYLARI:
- Ko'nikmalar: {criteria.get('skills', 'Korsatilmagan')}
- Tajriba: {criteria.get('experience', 'Korsatilmagan')}
- Ta'lim: {criteria.get('education', 'Korsatilmagan')}
- Tillar: {criteria.get('languages', 'Korsatilmagan')}
- Qo'shimcha: {criteria.get('additional', 'Korsatilmagan')}

MAXSUS TAHLIL TALABI:
{custom_prompt}

Quyidagi JSON formatda batafsil tahlil bering:
{{
    "compatibility_score": 85,
    "detailed_analysis": {{
        "skills_match": 90,
        "experience_match": 80,
        "education_match": 85,
        "overall_fit": "Yaxshi mos keladi"
    }},
    "strengths": [
        "Tegishli texnik ko'nikmalar",
        "Yetarli ish tajribasi",
        "Mos keluvchi ta'lim darajasi"
    ],
    "weaknesses": [
        "Ba'zi ko'nikmalar etishmaydi",
        "Maxsus loyiha tajribasi kam"
    ],
    "rejection_reasons": [
        "Minimal tajriba talabi bajarilmagan",
        "Asosiy ko'nikmalar yo'q"
    ],
    "recommendations": [
        "Qo'shimcha training olish",
        "Portfolio yaratish"
    ],
    "decision_rationale": "Nomzod asosiy talablarga javob beradi..."
}}

BAHOLASH MEZONLARI:
- 90-100%: A'lo nomzod
- 75-89%: Yaxshi nomzod  
- 60-74%: Qabul qilinadigan nomzod
- 0-59%: Rad etiladigan nomzod

Objektiv va adolatli baholang!
"""
        
        return prompt
    
    def _format_resume_for_analysis(self, resume: Dict) -> str:
        """Resume ma'lumotlarini tahlil uchun formatlash"""
        formatted = []
        
        # Shaxsiy ma'lumotlar
        personal = resume.get('personal_info', {})
        if personal.get('name'):
            formatted.append(f"Ism: {personal['name']}")
        
        # Ta'lim
        education = resume.get('education', [])
        if education:
            formatted.append("TA'LIM:")
            for edu in education[:3]:  # Birinchi 3 ta
                degree = edu.get('degree', 'N/A')
                field = edu.get('field', 'N/A')
                institution = edu.get('institution', 'N/A')
                year = edu.get('graduation_year', 'N/A')
                formatted.append(f"- {degree} {field} ({institution}, {year})")
        
        # Ish tajribasi
        experience = resume.get('work_experience', [])
        if experience:
            formatted.append("ISH TAJRIBASI:")
            for exp in experience[:3]:  # Birinchi 3 ta
                position = exp.get('position', 'N/A')
                company = exp.get('company', 'N/A')
                duration = exp.get('duration', 'N/A')
                formatted.append(f"- {position} ({company}, {duration})")
        
        # Ko'nikmalar
        skills = resume.get('skills', {})
        if skills.get('technical_skills'):
            tech_skills = ', '.join(skills['technical_skills'][:10])  # Birinchi 10 ta
            formatted.append(f"TEXNIK KO'NIKMALAR: {tech_skills}")
        
        if skills.get('languages'):
            languages = ', '.join(skills['languages'][:5])  # Birinchi 5 ta
            formatted.append(f"TILLAR: {languages}")
        
        # Loyihalar
        projects = resume.get('projects', [])
        if projects:
            formatted.append("LOYIHALAR:")
            for proj in projects[:2]:  # Birinchi 2 ta
                name = proj.get('name', 'N/A')
                tech = ', '.join(proj.get('technologies', [])[:3])
                formatted.append(f"- {name} ({tech})")
        
        return '\n'.join(formatted)
    
    def _parse_compatibility_result(self, response_text: str) -> Dict[str, Any]:
        """AI javobini parse qilish"""
        try:
            # JSON qismini topish
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group())
                
                # Asosiy fieldlarni tekshirish
                if 'compatibility_score' not in result:
                    result['compatibility_score'] = 50
                
                return result
            else:
                logger.warning("No JSON found in compatibility analysis response")
                return self._get_default_compatibility_result()
                
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error in compatibility analysis: {e}")
            return self._get_default_compatibility_result()
    
    def _validate_compatibility_result(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """Compatibility analysis natijasini validatsiya qilish"""
        # Score ni 0-100 oralig'ida cheklash
        score = result.get('compatibility_score', 0)
        if not isinstance(score, (int, float)):
            score = 0
        result['compatibility_score'] = max(0, min(100, int(score)))
        
        # Asosiy fieldlarni tekshirish va to'ldirish
        if 'strengths' not in result or not isinstance(result['strengths'], list):
            result['strengths'] = []
        
        if 'weaknesses' not in result or not isinstance(result['weaknesses'], list):
            result['weaknesses'] = []
        
        if 'rejection_reasons' not in result or not isinstance(result['rejection_reasons'], list):
            result['rejection_reasons'] = []
        
        # Agar score 60 dan past bo'lsa, rejection_reasons bo'lishi kerak
        if result['compatibility_score'] < Config.MIN_MATCH_PERCENTAGE:
            if not result['rejection_reasons']:
                result['rejection_reasons'] = [
                    "Minimal talablarga javob bermaydi",
                    "Ko'nikmalar yetarli emas",
                    "Tajriba darajasi past"
                ]
        
        return result
    
    def _get_default_compatibility_result(self) -> Dict[str, Any]:
        """Default compatibility result"""
        return {
            "compatibility_score": 0,
            "detailed_analysis": {
                "skills_match": 0,
                "experience_match": 0,
                "education_match": 0,
                "overall_fit": "Ma'lumot yetarli emas"
            },
            "strengths": [],
            "weaknesses": ["Tahlil qilishda xatolik"],
            "rejection_reasons": ["Technical xatolik"],
            "recommendations": ["Qayta urinib ko'ring"],
            "decision_rationale": "Tahlil muvaffaqiyatsiz tugadi"
        }
    
    # ===========================================
    # INTERVIEW EVALUATION
    # ===========================================
    
    async def evaluate_interview_answers(self, answers: List[Dict[str, str]], 
                                       ai_criteria: Dict[str, Any], 
                                       ai_prompt: str, 
                                       vacancy: Dict[str, Any]) -> Dict[str, Any]:
        """Suhbat javoblarini AI orqali baholash"""
        try:
            # Interview baholash prompt yaratish
            evaluation_prompt = self._create_interview_evaluation_prompt(
                answers, ai_criteria, ai_prompt, vacancy
            )
            
            response = await self._make_ai_request(
                system_message="Siz professional HR suhbat o'tkazuvchisigiz. Nomzod javoblarini baholaysiz.",
                user_message=evaluation_prompt,
                max_tokens=1200
            )
            
            # Natijani parse qilish
            evaluation_result = self._parse_interview_result(response, len(answers))
            
            # Validatsiya
            evaluation_result = self._validate_interview_result(evaluation_result)
            
            logger.info(f"Interview evaluation completed: {evaluation_result.get('overall_score', 0)}%")
            return evaluation_result
            
        except Exception as e:
            logger.error(f"Interview evaluation error: {e}")
            return self._get_default_interview_result(len(answers))
    
    def _create_interview_evaluation_prompt(self, answers: List[Dict[str, str]], 
                                          criteria: Dict[str, Any], 
                                          custom_prompt: str, 
                                          vacancy: Dict[str, Any]) -> str:
        """Interview baholash uchun prompt yaratish"""
        
        # Savol-javoblarni formatlash
        qa_formatted = []
        for i, answer_obj in enumerate(answers, 1):
            question = answer_obj.get('question', 'Savol yo\'q')
            answer = answer_obj.get('answer', 'Javob yo\'q')
            qa_formatted.append(f"SAVOL {i}: {question}\nJAVOB {i}: {answer}\n")
        
        qa_text = "\n".join(qa_formatted)
        
        prompt = f"""
SUHBAT JAVOBLARINI PROFESSIONAL BAHOLASH:

VAKANSIYA MA'LUMOTLARI:
- Nomi: {vacancy.get('title', 'N/A')}
- Tavsif: {vacancy.get('description', 'N/A')[:200]}...
- Talablar: {vacancy.get('requirements', 'N/A')[:200]}...

BAHOLASH KRITERIYLARI:
- Ko'nikmalar: {criteria.get('skills', 'N/A')}
- Tajriba: {criteria.get('experience', 'N/A')}
- Ta'lim: {criteria.get('education', 'N/A')}

MAXSUS BAHOLASH TALABI:
{custom_prompt}

NOMZOD JAVOBI:
{qa_text}

Quyidagi JSON formatda batafsil baholang:
{{
    "overall_score": 85,
    "individual_scores": [80, 90, 85],
    "detailed_evaluation": {{
        "communication_skills": 90,
        "technical_knowledge": 80,
        "problem_solving": 85,
        "cultural_fit": 88
    }},
    "positive_aspects": [
        "Aniq va tushunarli javoblar",
        "Yaxshi texnik bilim",
        "Ijodiy yondashish"
    ],
    "improvement_areas": [
        "Batafsil misollar kam",
        "Ba'zi texnik tafsilotlar etishmaydi"
    ],
    "overall_impression": "Nomzod umumiy taassurot...",
    "recommendation": "Qabul qilish/Rad etish",
    "passed": true
}}

BAHOLASH MEZONLARI:
- 85-100%: A'lo nomzod
- 70-84%: Yaxshi nomzod (minimal chegara)
- 55-69%: O'rtacha nomzod
- 0-54%: Rad etiladigan nomzod

Har bir javobni alohida baholang va umumiy xulosaga keling!
"""
        
        return prompt
    
    def _parse_interview_result(self, response_text: str, questions_count: int) -> Dict[str, Any]:
        """Interview baholash natijasini parse qilish"""
        try:
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                result = json.loads(json_match.group())
                return result
            else:
                logger.warning("No JSON found in interview evaluation response")
                return self._get_default_interview_result(questions_count)
                
        except json.JSONDecodeError as e:
            logger.error(f"Interview evaluation JSON parsing error: {e}")
            return self._get_default_interview_result(questions_count)
    
    def _validate_interview_result(self, result: Dict[str, Any]) -> Dict[str, Any]:
        """Interview baholash natijasini validatsiya qilish"""
        # Overall score ni tekshirish
        score = result.get('overall_score', 0)
        if not isinstance(score, (int, float)):
            score = 0
        result['overall_score'] = max(0, min(100, int(score)))
        
        # Passed statusini belgilash
        result['passed'] = result['overall_score'] >= Config.MIN_INTERVIEW_SCORE
        
        # Asosiy fieldlarni tekshirish
        for field in ['positive_aspects', 'improvement_areas']:
            if field not in result or not isinstance(result[field], list):
                result[field] = []
        
        # Individual scores ni tekshirish
        if 'individual_scores' not in result or not isinstance(result['individual_scores'], list):
            result['individual_scores'] = [result['overall_score']]
        
        return result
    
    def _get_default_interview_result(self, questions_count: int) -> Dict[str, Any]:
        """Default interview result"""
        return {
            "overall_score": 0,
            "individual_scores": [0] * questions_count,
            "detailed_evaluation": {
                "communication_skills": 0,
                "technical_knowledge": 0,
                "problem_solving": 0,
                "cultural_fit": 0
            },
            "positive_aspects": [],
            "improvement_areas": ["Baholash amalga oshirilmadi"],
            "overall_impression": "Technical xatolik tufayli baholanmadi",
            "recommendation": "Qayta suhbat o'tkazish",
            "passed": False
        }
    
    # ===========================================
    # UTILITY METHODS
    # ===========================================
    
    @retry(
        stop=stop_after_attempt(3),
        wait=wait_exponential(multiplier=1, min=4, max=10),
        retry=retry_if_exception_type((openai.RateLimitError, openai.APITimeoutError))
    )
    async def _make_ai_request(self, system_message: str, user_message: str, 
                             max_tokens: int = None) -> str:
        """AI ga so'rov yuborish - retry logic bilan"""
        try:
            # Rate limiting
            await self._check_rate_limit()
            
            response = await self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": system_message},
                    {"role": "user", "content": user_message}
                ],
                temperature=self.temperature,
                max_tokens=max_tokens or self.max_tokens
            )
            
            result = response.choices[0].message.content
            
            # Request counterni oshirish
            self.request_count += 1
            self.last_request_time = time.time()
            
            return result
            
        except openai.RateLimitError as e:
            logger.warning(f"Rate limit reached: {e}")
            raise
        except openai.APITimeoutError as e:
            logger.warning(f"API timeout: {e}")
            raise  
        except Exception as e:
            logger.error(f"AI request error: {e}")
            raise
    
    async def _check_rate_limit(self):
        """Rate limiting tekshirish"""
        current_time = time.time()
        
        # Yangi oyna boshlash
        if current_time - self.request_window_start > 60:
            self.request_count = 0
            self.request_window_start = current_time
        
        # Rate limit tekshirish
        if self.request_count >= Config.OPENAI_MAX_REQUESTS_PER_MINUTE:
            sleep_time = 60 - (current_time - self.request_window_start)
            if sleep_time > 0:
                logger.info(f"Rate limit reached, waiting {sleep_time:.1f} seconds")
                await asyncio.sleep(sleep_time)
                self.request_count = 0
                self.request_window_start = time.time()
        
        # So'rovlar orasida minimal kutish
        if current_time - self.last_request_time < 1:
            await asyncio.sleep(1 - (current_time - self.last_request_time))


# ===========================================
# HELPER FUNCTIONS
# ===========================================

def validate_file_format(filename: str) -> bool:
    """Fayl formatini tekshirish"""
    if not filename:
        return False
    
    file_extension = Path(filename).suffix.lower()
    return file_extension in Config.ALLOWED_FILE_EXTENSIONS


def get_file_size_mb(file_path: str) -> float:
    """Fayl hajmini MB da qaytarish"""
    try:
        size_bytes = os.path.getsize(file_path)
        return round(size_bytes / (1024 * 1024), 2)
    except Exception:
        return 0.0


def create_safe_filename(original_filename: str, user_id: int) -> str:
    """Xavfsiz fayl nomi yaratish"""
    if not original_filename:
        return f"{user_id}_{int(time.time())}.pdf"
    
    # Faqat xavfsiz belgilarni qoldirish
    safe_name = re.sub(r'[^a-zA-Z0-9._-]', '_', original_filename)
    
    # Maksimal uzunlik
    name, extension = os.path.splitext(safe_name)
    if len(name) > 50:
        name = name[:50]
    
    # Vaqt va user_id qo'shish
    timestamp = int(time.time())
    return f"{user_id}_{timestamp}_{name}{extension}"


async def test_ai_service(api_key: str) -> bool:
    """AI Service ni test qilish"""
    try:
        ai_service = AIService(api_key)
        return await ai_service.test_connection()
    except Exception as e:
        logger.error(f"AI Service test failed: {e}")
        return False


# Factory function
def create_ai_service() -> AIService:
    """AI Service yaratish - factory pattern"""
    return AIService(
        api_key=Config.OPENAI_API_KEY,
        model=Config.OPENAI_MODEL,
        max_tokens=Config.OPENAI_MAX_TOKENS,
        temperature=Config.OPENAI_TEMPERATURE
    )